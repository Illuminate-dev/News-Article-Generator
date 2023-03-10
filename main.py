import os
import re
import datetime

import openai
import requests
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from dotenv import load_dotenv
from serpapi import GoogleSearch

load_dotenv()

openai.api_key = os.getenv("OPENAI_API_KEY")
serpapi_key = os.getenv('SERPAPI_KEY')

TEMPERATURE=1

DAVINCI_MODEL="text-davinci-003"
# WHEN USING CUSTOM MODEL, PROMPTS HAVE TO END WITH "\n\n###\n\n"
# DONT START WITH, write a news article about
# include stop=[" END"]
CUSTOM_MODEL="davinci:ft-personal-2023-01-06-23-00-19"

TRAINING_DATE = datetime.date(2021, 6, 1)

CUSTOM_PROMPT_1="{}.\n"
CUSTOM_PROMPT_2="Include the following information:\n{}\n\nPeople/Organizations:\n{}"
CUSTOM_PROMPT_3="Date: {}"
CUSTOM_PROMPT_END = "\n\n###\n\n"

# PROMPT CONFIG
TOPIC_MESSAGE="Write a news article about {}.\n"
NOTES_MESSAGE="The article should include the following information:\n{}\n"
CAPTION_MESSAGE="Generate two captions for the article which clearly depict a scene. Surround the captions with brackets like this: [caption here] and choose appropriate locations within the article to place them. "
REQUIREMENTS_MESSAGE="Include a title at the start surrounded by parentheses like this: (title here). This is a lengthy article and must be at least {} words long."

IMAGE_PROMPT="An award winning professional photo with the caption of \"{}\", realistic lighting, photojournalism from The New York Times"

def download_image(url, name):
    img = requests.get(url).content
    path = 'png/' + name + '.png'
    if not os.path.isdir('png'):
        os.mkdir('png')
    with open(path, "wb") as f:
       f.write(img) 
    return path

def create_custom_prompt():
    prompt = ""
    topic = input("Write a news article about: ")
    prompt += CUSTOM_PROMPT_1.format(topic)

    inp = input("mm/yyyy of event: (put nothing if no event)").split('/')
    date = datetime.date(int(inp[1]),int(inp[0]),1)
    if date >= TRAINING_DATE:
        notes = []
        while True:
            line = input("Notes of info to add (empty for none): ")
            if line:
                notes.append(line)
            elif len(notes)!=0:
                break
        qoutes = []
        while True:
            line = input("People/Qoutes to add (empty for none): ")
            if line:
                qoutes.append(line)
            elif len(qoutes) != 0:
                break
        prompt += CUSTOM_PROMPT_2.format('\n'.join(notes), '\n'.join(qoutes))
    else:
        prompt += CUSTOM_PROMPT_3.format(date.strftime("%B %Y"))
    prompt += CUSTOM_PROMPT_END
    return prompt
    


def create_prompt():

    topic = input("What to write about: ")
    notes = []
    while True:
        line = input("Optional notes of info to add (empty for none): ")
        if line:
            notes.append(line)
        else:
            break

    prompt = TOPIC_MESSAGE.format(topic)

    # append notes if notes is not empty
    prompt += NOTES_MESSAGE.format('\n'.join(notes)) if notes.count != 0 else "" 
        
    prompt += CAPTION_MESSAGE

    word_count = input("How many words? ")

    prompt += REQUIREMENTS_MESSAGE.format(word_count if len(word_count) != 0 else 400)

    return prompt

def generate(prompt, model, images=True):
    response = openai.Completion.create(engine=model, prompt=prompt, max_tokens=3500 if model.startswith('text') else 2049-(len(prompt)//4)-15, temperature=TEMPERATURE)

    text = response.get("choices")[0]["text"]

    print(text)

    title = re.findall('\((.*?)\)', text)[0]
    print(f'Title: {title}')

    captions = re.findall('\[(.*?)\]', text)
    paths = []

    if images:
        for caption in captions:
            image = openai.Image.create(prompt=IMAGE_PROMPT.format(caption), n=1, size="1024x1024")
            url = image['data'][0]['url']
            filename = '_'.join(caption.split(' '))[0:14]
            path = download_image(url, filename)
            paths.append(path)
    
    return text, title, captions, paths

def search_image(query):
    search = GoogleSearch({
        'q': query,
        'tbm': 'isch',
        'api_key': serpapi_key
    })
    return search.get_dict()['images_results'][0]['original']


def save_doc(text, title, captions, paths, filename):
    pg_counter = 0

    document = Document()

    doc_title = document.add_heading(title)
    doc_title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    pg_counter += 1

    text_style = document.styles.add_style("text_01", WD_STYLE_TYPE.PARAGRAPH)
    text_style.font.name = "Times New Roman"
    text_style.font.size = Pt(13)

    caption_style = document.styles.add_style("caption_01", WD_STYLE_TYPE.PARAGRAPH)
    caption_style.font.size = Pt(9)
    caption_style.font.name = "Times New Roman"

    pathindex = 0
    for line in text.split('\n'):
        if len(line) == 0 or line.startswith("(") or line.startswith(" ("):
            continue
        elif line.startswith("["):

            document.add_picture(paths[pathindex], Inches(5))   

            pg_counter += 1

            document.paragraphs[pg_counter-1].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER 
            
            paragraph = document.add_paragraph(captions[pathindex])

            pg_counter+= 1

            paragraph.style = caption_style
            paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

            pathindex += 1
        else:
            doc_pg = document.add_paragraph(line) 
            doc_pg.style = text_style
            pg_counter += 1
        

        document.save(filename)

def main():
    print(
"""[0]: DALL-E create images
[1]: SerpAPI find images
[2]: provide images""")
    choice = int(input("Enter choice: "))
    if choice == 0:
        prompt = create_prompt()
        text, title, captions, paths = generate(prompt=prompt, model=DAVINCI_MODEL)
        filename = '_'.join(title.split(' '))[0:14] + '.docx'
        save_doc(text, title, captions, paths, filename)
    elif choice == 1:
        prompt = create_custom_prompt()
        text, title, captions, paths = generate(prompt, model=CUSTOM_MODEL, images=False)
        text = text.replace("END", '')
        for caption in captions:
            link = search_image(caption)
            filename = '_'.join(caption.split(' '))[0:14]
            path = download_image(link, filename)
            paths.append(path)
        filename = '_.'.join(title.split(' '))[0:14] + '.docx'
        save_doc(text, title, captions, paths, filename)
    elif choice == 2:
        prompt = create_custom_prompt()
        text, title, captions, paths = generate(prompt, model=CUSTOM_MODEL, images=False)
        text = text[:-4].replace("END", '')
        newtext = ''
        for line in text.splitlines():
            if not line.startswith('['):
                newtext += line + '\n'
        filename = '_.'.join(title.split(' '))[0:14] + '.docx'
        save_doc(newtext, title, captions, paths, filename)

if __name__ == '__main__':

    main()